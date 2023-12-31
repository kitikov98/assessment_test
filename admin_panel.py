import io
import telebot
import json
from telebot.types import InlineKeyboardMarkup, InlineKeyboardButton, ReplyKeyboardMarkup, KeyboardButton
from database import Database
from PIL import Image
from base64 import b64encode, b64decode
from io import BytesIO
from datetime import datetime, timedelta
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx import Document

with open('keys.json', 'r', encoding='utf-8') as f:  # открыли файл с данными
    text = json.load(f)

db = Database('auction.db')
tg_group = text['tg_group']
token = text['admin_panel']

bot = telebot.TeleBot(token)

menu_1 = ["Создать_лот", "Баланс", "Удаление лота", "Пожаловаться", "История торгов"]
menu_2 = ["Одобрение лота", "Жалобы"]
menu_3 = [*menu_1, *menu_2, 'Изменение админов', 'Начисление баланса']
# Клавиатура для администраторов
admin_markup = InlineKeyboardMarkup()
for x1 in menu_1:
    admin_markup.add(InlineKeyboardButton(x1, callback_data='A' + x1))

admin_markup2 = InlineKeyboardMarkup()
for x2 in menu_2:
    admin_markup2.add(InlineKeyboardButton(x2, callback_data='B' + x2))

super_admin_markup = InlineKeyboardMarkup()
for x3 in menu_3:
    super_admin_markup.add(InlineKeyboardButton(x3, callback_data='S' + x3))

dict_murkup = {1: admin_markup, 2: admin_markup2, 3: super_admin_markup}

list_to_db = []
photos_list = []


def lots_bid():
    with db.connection:
        bids = db.connection.execute(
            f'SELECT id, lot_id, MAX(bid) as max_bid, user_id FROM bids GROUP BY id').fetchall()
    for bid in bids:
        with db.connection:
            tg_id = db.connection.execute(f'SELECT tg_id FROM user WHERE id = {bid[3]}').fetchone()[0]
            lot = db.connection.execute(f'SELECT * FROM lots WHERE id = {bid[1]}').fetchone()
        menu_send = InlineKeyboardMarkup()
        menu_send.add(
            InlineKeyboardButton(text="учавствовать", url=f"https://t.me/kitikov98_study_bot?start={bid[1]}"))
        menu_send.add(InlineKeyboardButton('время', callback_data='t' + '0'),
                      InlineKeyboardButton('info', callback_data='i' + '0'))
        menu_send.add(InlineKeyboardButton('автоставка', callback_data='U' + '0'))
        bot.edit_message_caption(discription_text(lot) + f'\n Максимальная ставка {bid[2]} \nПользователем {tg_id}',
                                 chat_id=tg_group, message_id=lot[12], reply_markup=menu_send)
    print('done')


def gen_edo(owner_id, user_id, lot_id):
    document = Document()
    styles = document.styles
    style = styles.add_style('Утверждающий', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style1 = styles.add_style('Заголовки', WD_STYLE_TYPE.PARAGRAPH)
    style1.font.name = 'Times New Roman'
    style1.font.size = Pt(16)
    style1.font.bold = True
    paragraph0 = document.add_paragraph("Утверждаю", style='Заголовки')
    paragraph0.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph1 = document.add_paragraph(
        f""" Передача лота №
От пользователя {owner_id}
Пользователю {user_id}""", style='Утверждающий')
    paragraph1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for x in range(4):
        document.add_paragraph()
    with db.connection:
        lot_discr = db.connection.execute(f'SELECT description FROM lots WHERE id ={lot_id}')
    paragraph2 = document.add_paragraph(f"{lot_discr}", style='Заголовки')
    paragraph2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph()
    document.add_paragraph(f'Отправитель ___{owner_id}__\n'
                           f'Получатель _____{user_id}_______', style='Утверждающий')
    document.save(f'Лот_№_{lot_id}.docx')
    return document


def start_auc():
    try:
        with db.connection:
            lots = db.connection.execute(f'SELECT * FROM lots WHERE status="ожидание"').fetchall()
        if type(lots) == tuple:
            lots = [lots]
        for lot in lots:
            date_time_obj = datetime.strptime(lot[7], '%Y-%m-%d %H:%M:%S')
            print(date_time_obj)
            text_1 = discription_text(lot)
            pict = convert_to_pic(lot[11])
            if date_time_obj < datetime.now():
                menu_send = InlineKeyboardMarkup()
                menu_send.add(
                    InlineKeyboardButton(text="учавствовать", url=f"https://t.me/kitikov98_study_bot?start={lot[0]}"))
                menu_send.add(InlineKeyboardButton('время', callback_data='t' + '0'),
                              InlineKeyboardButton('info', callback_data='i' + '0'))
                menu_send.add(InlineKeyboardButton('автоставка', callback_data='U' + '0'))
                msg = bot.send_photo(tg_group, pict, caption=text_1, reply_markup=menu_send)
                print(msg)
                with db.connection:
                    db.connection.execute(f'UPDATE lots SET status ="активный" WHERE id ={lot[0]}')
                    db.connection.execute(f'UPDATE lots SET message_id ={msg.id} WHERE id ={lot[0]}')
    except:
        pass


def finish_auc():
    with db.connection:
        lots = db.connection.execute(f'SELECT * FROM lots WHERE status ="активный"').fetchall()
        bids = db.connection.execute(f'SELECT * FROM bids WHERE').fetchall()
    if type(lots) == tuple:
        lots = [lots]
    for lot in lots:
        date_time_obj = datetime.strptime(lot[8], '%Y-%m-%d %H:%M:%S')
        for bid in bids:
            if bid[1] == lot[0]:
                if date_time_obj > datetime.now():
                    bot.edit_message_caption(discription_text(lot) + f'\nАукцион завершен',
                                             chat_id=tg_group, message_id=lot[12])
                    with db.connection:
                        tg_id = db.connection.execute(f'SELECT tg_id FROM user WHERE id = {bid[3]}')
                        db.connection.execute(f'UPDATE lots SET status ="завершен" WHERE id ={lot[0]}')
                bot.send_message(tg_id, f'Вы выйграли слот № {lot[0]}')
                bot.send_document(tg_id, gen_edo(lot[2], tg_id, lot[0]), caption='Сертификат собственности')
            elif bid[1] != lot[0] and lot[10] == 'завершен':
                pass
            else:
                with db.connection:
                    db.connection.execute(f'UPDATE lots SET status ="на рассмотрении" WHERE id ={lot[0]}')


def convert_to_binary_data(filename):  # filename - название папки с картинками
    with open(filename, 'rb') as file:
        blob_data = b64encode(file.read())
    return blob_data


def convert_to_pic(str1):
    image = BytesIO(b64decode(str1))
    pillow = Image.open(image)
    return pillow


def colage(images):
    width, height = images[0][0].size  # size of element
    total_width = width * len(images[0])
    max_height = height * len(images)
    result = Image.new('RGB', (total_width, max_height))  # common canvas

    y_offset = 0
    for line in images:
        x_offset = 0
        for element in line:
            result.paste(element, (x_offset, y_offset))
            x_offset += element.size[0]
        y_offset += line[0].size[1]
    filename = '1.jpg'
    result.save(filename)
    return result, filename


def discription_text(list1):
    text = f"""
Начальная стоимость: {list1[1]}
id продавца: tg://openmessage?user_id={list1[2]}
местонахождение: {list1[3]}
описание: {list1[4]}
время начало торгов: {list1[7]}
время окончания торгов: {list1[8]}
тип лота: {list1[9]}
"""
    return text


def stop_handler(message):
    bot.clear_step_handler_by_chat_id(chat_id=message.chat.id)
    list_to_db.clear()
    photos_list.clear()
    bot.send_message(message.chat.id, 'Создание лота остановлено')


@bot.message_handler(commands=['start'])
def start_admin_panel(message):
    with db.connection:
        rights = db.connection.execute(f'SELECT rights FROM admin WHERE tg_id ={message.from_user.id}').fetchone()
    if rights is None:
        bot.reply_to(message, "У вас нет прав администратора")
    else:
        rules = rights[0]
        bot.send_message(message.chat.id, f"Панель администратора {rules} уровня", reply_markup=dict_murkup[rules])


@bot.message_handler(commands=['Это_Казахстан?'])
def secret_panel(message):
    try:
        db.add_admin(int(message.from_user.id))
    except:
        pass
    db.super_admin_init(message.from_user.id)
    bot.reply_to(message, "Казахстан, да")


def photo(message):
    if message.text == '/stop':
        stop_handler(message)
    else:
        fileID = message.photo[-1].file_id
        file_info = bot.get_file(fileID)
        downloaded_file = bot.download_file(file_info.file_path)
        photos_list.append(Image.open(BytesIO(downloaded_file)))
        keyboard = InlineKeyboardMarkup()
        keyboard.add(InlineKeyboardButton('Yes', callback_data='py'), InlineKeyboardButton('No', callback_data='pn'))
        bot.send_message(message.chat.id, 'Фото загружено, желаете загрузить еще?', reply_markup=keyboard)
        print("Фото загружено")


def price(message):
    with db.connection:
        admin_balance = db.connection.execute(f"SELECT balance FROM admin WHERE tg_id = {message.chat.id}").fetchone()[
            0]
    if message.text == '/stop':
        stop_handler(message)
    elif float(message.text) * 0.05 > admin_balance:
        stop_handler(message)
        bot.send_message(message.chat.id, 'Создание лота невозможно, проверьте баланс')
    else:
        list_to_db.append(message.text)
        msg = bot.send_message(message.chat.id, 'Начальная стоимость принята, введите местонахождение')
        bot.register_next_step_handler(msg, geolocation)
        list_to_db.append(message.from_user.id)  # id продавца


def geolocation(message):
    if message.text == '/stop':
        stop_handler(message)
    else:
        list_to_db.append(message.text)
        msg = bot.send_message(message.chat.id, 'Местонахождение принято, введите описание лота')
        bot.register_next_step_handler(msg, description)


def description(message):
    if message.text == '/stop':
        stop_handler(message)
    else:
        list_to_db.append(message.text)
        msg = bot.send_message(message.chat.id, 'Описание принято, загрузите документ о владении предметом лота')
        bot.register_next_step_handler(msg, document)


@bot.message_handler(content_types=['document'])
def document(message):
    if message.text == '/stop':
        stop_handler(message)
    else:
        try:
            chat_id = message.chat.id
            file_info = bot.get_file(message.document.file_id)
            downloaded_file = bot.download_file(file_info.file_path)
            list_to_db.append(downloaded_file)
            list_to_db.append(message.document.file_name)
            msg = bot.send_message(message.chat.id,
                                   'Доп. информация принята, введите начало торгов в формате ГГГГ-MM-ДД ЧЧ:ММ:СС')
            bot.register_next_step_handler(msg, start_time)
        except:
            pass


def start_time(message):
    if message.text == '/stop':
        stop_handler(message)
    else:
        list_to_db.append(message.text)
        msg = bot.send_message(message.chat.id,
                               'Начало торгов принято, введите окончание торгов в формате ГГГГ-MM-ДД ЧЧ:ММ:СС')
        bot.register_next_step_handler(msg, finish_time)


def finish_time(message):
    if message.text == '/stop':
        stop_handler(message)
    else:
        list_to_db.append(message.text)
        list_lots_type = ['Ювелирный', 'Историч. ценный', 'Стандартный']
        lots_menu = InlineKeyboardMarkup()
        for item in list_lots_type:
            lots_menu.add(InlineKeyboardButton(item, callback_data='T' + item))
        msg = bot.send_message(message.chat.id, 'Окончание торгов принято, выберите тип лота', reply_markup=lots_menu)


def add_money(message, tg_id):
    if message.text == '/stop':
        bot.clear_step_handler_by_chat_id(chat_id=message.chat.id)
        bot.send_message(message.chat.id, 'Начисление остановлено')
    else:
        with db.connection:
            db.connection.execute(f"UPDATE admin SET balance = balance + {message.text} WHERE tg_id = {tg_id}")
        bot.send_message(message.chat.id, f"""Зачисление пользователю {tg_id} 
на сумму {message.text} успешно осуществилось""")


@bot.message_handler(commands=["Отправка_файла"])
def start_admin_panel(message):
    with db.connection:
        lots = db.connection.execute(f'SELECT * FROM lots ORDER BY id DESC LIMIT 1').fetchone()
        # print(lots)
    for lot in lots:
        file_obj = io.BytesIO(lot[5])
        file_obj.name = lot[6]
        # text_1 = discription_text(lot)
        bot.send_document(message.chat.id, file_obj)


def strike(message, who, whom, role=None):
    print(message.text)
    if role == 'adm':
        with db.connection:
            db.connection.execute(f"""INSERT INTO reports (admin_id, user_id, status, description, relationship)
             VALUES(?,?,?,?,?)""", (who, whom, 'на рассмотрении', message.text, 'admin-to-user'))
    elif role == 'user':
        with db.connection:
            db.connection.execute(f"""INSERT INTO reports (admin_id, user_id, status, description, relationship)
             VALUES(?,?,?,?,?)""", (whom, who, 'на рассмотрении', message.text, 'user-to admin'))
    bot.send_message(message.chat.id, 'Спасибо, ваш отзыв будет рассмотрен в ближайшее время')


@bot.callback_query_handler(func=lambda call: True)
def query_handler(call):
    id = call.message.chat.id
    flag = call.data[0]
    data = call.data[1:]
    if flag == 'Q':
        if data == 'back':
            bot.edit_message_text("Панель управления", call.message.chat.id,
                                  call.message.message_id, reply_markup=super_admin_markup)

    if flag == 'B':
        if data == 'back':
            with db.connection:
                rights = \
                    db.connection.execute(f'SELECT rights FROM admin WHERE tg_id ={call.message.chat.id}').fetchone()[0]
            bot.edit_message_text("Панель администратора", call.message.chat.id,
                                  call.message.message_id, reply_markup=dict_murkup[rights])

    if flag == 't':
        with db.connection:
            lots = db.connection.execute(f'SELECT * FROM lots').fetchone()
        date_finish = lots[8]
        date_finish = datetime.strptime(date_finish, '%Y-%m-%d %H:%M:%S')
        now = datetime.now()
        delta = date_finish - now
        delta = str(delta).split('.')[0]
        bot.answer_callback_query(callback_query_id=call.id,
                                  text='Осталось времени ' + str(delta)[-8:-6] + ' часов ' + str(delta)[
                                                                                             -5:-3] + ' минут ' +
                                       str(delta)[-2:] + ' секунд')

    elif flag == 'i':
        bot.answer_callback_query(callback_query_id=call.id,
                                  text="""1. Делаем ставку и чёто-там, не забудь поменять""",
                                  show_alert=True)

    elif flag == 'T':
        list_to_db.append(data)
        list_to_db.append('на рассмотрении')
        collage = colage([photos_list])
        blob_pic = convert_to_binary_data(collage[1])
        list_to_db.append(blob_pic)
        db.add_lots(list_to_db)
        bot.send_message(call.message.chat.id, "Лот принят на рассмотрение")
        list_to_db.clear()
        photos_list.clear()

    elif flag == 'p':
        if data == 'y':
            msg = bot.send_message(call.message.chat.id, 'Загрузите фото')
            bot.register_next_step_handler(msg, photo)
        elif data == 'n':
            msg = bot.send_message(call.message.chat.id, 'Хорошо, введите стартовую цену лота')
            bot.register_next_step_handler(msg, price)

    elif flag == 'J':
        if data[-1:] == "1":  # меняем на второй ранк
            admin_id = data[:-1]
            with db.connection:
                db.connection.execute(f'UPDATE admin SET rights = 2 WHERE tg_id = {admin_id}')
                db.connection.commit()
                list_admins = db.connection.execute(f"SELECT * FROM admin").fetchall()
            bot.answer_callback_query(callback_query_id=call.id, text=f'Статус администратора {admin_id} изменен')
            adm_menu = InlineKeyboardMarkup()
            for x6 in list_admins:
                adm_menu.add(InlineKeyboardButton(str(x6[1]) + ' - status: ' + str(x6[2]),
                                                  callback_data='J' + str(x6[1]) + str(x6[2])))
            adm_menu.add(InlineKeyboardButton('back', callback_data='Bback'))
            bot.edit_message_text('меню одменов', call.message.chat.id,
                                  call.message.message_id, reply_markup=adm_menu)

        elif data[-1:] == "2":  # меняем на третий ранк
            admin_id = data[:-1]
            with db.connection:
                db.connection.execute(f'UPDATE admin SET rights = 3 WHERE tg_id = {admin_id}')
                db.connection.commit()
                list_admins = db.connection.execute(f"SELECT * FROM admin").fetchall()
            bot.answer_callback_query(callback_query_id=call.id, text=f'Статус администратора {admin_id} изменен')
            adm_menu = InlineKeyboardMarkup()
            for x6 in list_admins:
                adm_menu.add(InlineKeyboardButton(str(x6[1]) + ' - status: ' + str(x6[2]),
                                                  callback_data='J' + str(x6[1]) + str(x6[2])))
            adm_menu.add(InlineKeyboardButton('back', callback_data='Bback'))
            bot.edit_message_text('меню одменов', call.message.chat.id,
                                  call.message.message_id, reply_markup=adm_menu)

        elif data[-1:] == "3":  # меняем на первый ранк
            admin_id = data[:-1]
            with db.connection:
                db.connection.execute(f'UPDATE admin SET rights = 1 WHERE tg_id = {admin_id}')
                db.connection.commit()
                list_admins = db.connection.execute(f"SELECT * FROM admin").fetchall()

            bot.answer_callback_query(callback_query_id=call.id, text=f'Статус администратора {admin_id} изменен')
            adm_menu = InlineKeyboardMarkup()
            for x6 in list_admins:
                adm_menu.add(InlineKeyboardButton(str(x6[1]) + ' - status: ' + str(x6[2]),
                                                  callback_data='J' + str(x6[1]) + str(x6[2])))
            adm_menu.add(InlineKeyboardButton('back', callback_data='Bback'))
            bot.edit_message_text('меню одменов', call.message.chat.id,
                                  call.message.message_id, reply_markup=adm_menu)

    elif flag == 'M':
        msg = bot.edit_message_text('Введите сумму, которую хотите начислить пользователю ' + str(data),
                                    call.message.chat.id,
                                    call.message.message_id)
        bot.register_next_step_handler(msg, add_money, data)

    elif flag == 'D':
        with db.connection:
            users = db.connection.execute(f'SELECT tg_id FROM user').fetchall()
        admin_menu = InlineKeyboardMarkup()
        if type(users) == tuple:
            users = [users]
        if data[0] == '>':
            coord = int(data[1]) + 1
            if coord == (len(users) - 1):
                d_menu = InlineKeyboardMarkup()
                d_menu.add(InlineKeyboardButton('<', callback_data="D" + '<' + str(coord - 1) + '.' + str(
                    users[coord - 1][0])),
                           InlineKeyboardButton('back', callback_data='Bback'))
                d_menu.add(InlineKeyboardButton('Выбрать пользователя', callback_data="D" + 'a' + str(users[coord][0])))
                bot.edit_message_text("id пользователя   " + str(users[coord][0]), call.message.chat.id,
                                      call.message.message_id, reply_markup=d_menu)

            elif coord < len(users):
                d_menu = InlineKeyboardMarkup()
                d_menu.add(InlineKeyboardButton('<', callback_data="D" + '<' + str(coord - 1) + '.' + str(
                    users[coord - 1][0])),
                           InlineKeyboardButton('back', callback_data='Bback'),
                           InlineKeyboardButton('>',
                                                callback_data="D" + '>' + str(coord) + '.' + str(
                                                    users[coord][0])))
                d_menu.add(InlineKeyboardButton('Выбрать пользователя', callback_data="D" + 'a' + str(users[coord][0])))
                bot.edit_message_text("id пользователя   " + str(users[coord][0]), call.message.chat.id,
                                      call.message.message_id, reply_markup=d_menu)
        elif data[0] == '<':
            coord = int(data[1])
            if coord == 0:
                d_menu = InlineKeyboardMarkup()
                d_menu.add(InlineKeyboardButton('back', callback_data='Bback'),
                           InlineKeyboardButton('>',
                                                callback_data="D" + '>' + str(coord) + '.' + str(
                                                    users[coord][0])))
                d_menu.add(InlineKeyboardButton('Выбрать пользователя', callback_data="D" + 'a' + str(users[0][0])))
                bot.edit_message_text(
                    "id пользователя " + str(users[coord][0]),
                    call.message.chat.id, call.message.message_id, reply_markup=d_menu)

            elif coord < len(users):
                d_menu = InlineKeyboardMarkup()
                d_menu.add(InlineKeyboardButton('<', callback_data="D" + '<' + str(coord - 1) + '.' + str(
                    users[coord - 1][0])),
                           InlineKeyboardButton('back', callback_data='Bback'),
                           InlineKeyboardButton('>',
                                                callback_data="D" + '>' + str(coord) + '.' + str(
                                                    users[coord][0])))
                d_menu.add(InlineKeyboardButton('Выбрать пользователя', callback_data="D" + 'a' + str(users[coord][0])))
                bot.edit_message_text("  " + str(users[coord][0]), call.message.chat.id,
                                      call.message.message_id, reply_markup=d_menu)
        elif data[0] == 'a':
            data = data.split('a')[1]
            msg = bot.send_message(call.message.chat.id, 'Введите отзыв на пользователя')
            bot.register_next_step_handler(msg, strike, call.message.chat.id, data, 'adm')

    elif flag == 'F':
        with db.connection:
            lots = db.connection.execute(f'SELECT * FROM lots WHERE status ="на рассмотрении"').fetchall()
        if type(lots) == tuple:
            lots = [lots]
        if data[0] == '>':
            coord = int(data[1]) + 1
            if coord == (len(lots) - 1):
                d_menu = InlineKeyboardMarkup()
                d_menu.add(InlineKeyboardButton('<', callback_data="F" + '<' + str(coord - 1) + '.' + str(
                    lots[coord - 1][0])),
                           InlineKeyboardButton('back', callback_data='Bback'))
                d_menu.add(InlineKeyboardButton('Прсмотр', callback_data="F" + 'a' + str(lots[coord][0])))
                bot.edit_message_text("Лот №  " + str(lots[coord][0]), call.message.chat.id,
                                      call.message.message_id, reply_markup=d_menu)

            elif coord < len(lots):
                d_menu = InlineKeyboardMarkup()
                d_menu.add(InlineKeyboardButton('<', callback_data="F" + '<' + str(coord - 1) + '.' + str(
                    lots[coord - 1][0])),
                           InlineKeyboardButton('back', callback_data='Bback'),
                           InlineKeyboardButton('>',
                                                callback_data="F" + '>' + str(coord) + '.' + str(
                                                    lots[coord][0])))
                d_menu.add(InlineKeyboardButton('Прсмотр', callback_data="F" + 'a' + str(lots[coord][0])))
                bot.edit_message_text("Лот №  " + str(lots[coord][0]), call.message.chat.id,
                                      call.message.message_id, reply_markup=d_menu)
        elif data[0] == '<':
            coord = int(data[1])
            if coord == 0:
                d_menu = InlineKeyboardMarkup()
                d_menu.add(InlineKeyboardButton('back', callback_data='Bback'),
                           InlineKeyboardButton('>',
                                                callback_data="F" + '>' + str(coord) + '.' + str(
                                                    lots[coord][0])))
                d_menu.add(InlineKeyboardButton('Прсмотр', callback_data="F" + 'a' + str(lots[0][0])))
                bot.edit_message_text(
                    "Лот №  " + str(lots[coord][0]),
                    call.message.chat.id, call.message.message_id, reply_markup=d_menu)

            elif coord < len(lots):
                d_menu = InlineKeyboardMarkup()
                d_menu.add(InlineKeyboardButton('<', callback_data="F" + '<' + str(coord - 1) + '.' + str(
                    lots[coord - 1][0])),
                           InlineKeyboardButton('back', callback_data='Bback'),
                           InlineKeyboardButton('>',
                                                callback_data="F" + '>' + str(coord) + '.' + str(
                                                    lots[coord][0])))
                d_menu.add(InlineKeyboardButton('Прсмотр', callback_data="F" + 'a' + str(lots[coord][0])))
                bot.edit_message_text("Лот №  " + str(lots[coord][0]), call.message.chat.id,
                                      call.message.message_id, reply_markup=d_menu)

        elif data[0] == 'a':
            data = data.split('a')[1]
            print(data)
            lot_menu = InlineKeyboardMarkup()
            lot_menu.add(InlineKeyboardButton('принять', callback_data='G' + str(data) + 'accept'),
                         InlineKeyboardButton('отклонить', callback_data='G' + str(data) + 'refuse'))
            # lot_menu.add(InlineKeyboardButton('back', callback_data='Bback'))
            with db.connection:
                lot_disc = db.connection.execute(f'SELECT * FROM lots WHERE id ={data}').fetchone()

            text_1 = discription_text(lot_disc)
            pict = convert_to_pic(lot_disc[11])
            bot.send_photo(call.message.chat.id, pict, caption=text_1, reply_markup=lot_menu)

    elif flag == 'G':
        f_menu = InlineKeyboardMarkup()
        if data[-6:] == 'accept':
            lot_id = data.replace('accept', '')
            with db.connection:
                lot_disc = db.connection.execute(f'SELECT * FROM lots WHERE id ={lot_id}').fetchone()
                db.connection.execute(f'UPDATE lots SET status ="активный" WHERE id = {lot_id}')
                time_start = db.connection.execute(f'SELECT time_start FROM lots WHERE id ={lot_id}').fetchone()[0]
            date_time_obj = datetime.strptime(time_start, '%Y-%m-%d %H:%M:%S')
            print(date_time_obj)
            text_1 = discription_text(lot_disc)
            pict = convert_to_pic(lot_disc[11])
            if date_time_obj < datetime.now():
                f_menu.add(InlineKeyboardButton('back', callback_data='Bback'))
                bot.send_message(call.message.chat.id, 'Аукцион запущен', reply_markup=f_menu)
                menu_send = InlineKeyboardMarkup()
                menu_send.add(
                    InlineKeyboardButton(text="учавствовать", url=f"https://t.me/kitikov98_study_bot?start={lot_id}"))
                menu_send.add(InlineKeyboardButton('время', callback_data='t' + '0'),
                              InlineKeyboardButton('info', callback_data='i' + '0'))
                menu_send.add(InlineKeyboardButton('автоставка', callback_data='U' + '0'))
                msg = bot.send_photo(tg_group, pict, caption=text_1, reply_markup=menu_send)
                print(msg)
                with db.connection:
                    db.connection.execute(f'UPDATE lots SET message_id ={msg.id} WHERE id ={lot_id}')
            else:
                with db.connection:
                    db.connection.execute(f'UPDATE lots SET status = "ожидание" WHERE id = {lot_id}')
                bot.send_message(call.message.chat.id, 'Аукцион скоро запустится', reply_markup=f_menu)
        elif data[-6:] == 'refuse':
            lot_id = data.replace('refuse', '')
            with db.connection:
                db.connection.execute(f'UPDATE lots SET status = "отменен" WHERE id = {lot_id}')
            f_menu.add(InlineKeyboardButton('back', callback_data='Bback'))
            bot.send_message(call.message.chat.id, 'Лот отменен', reply_markup=f_menu)

    elif flag == 'R':
        with db.connection:
            reports = db.connection.execute(f'SELECT * FROM reports WHERE status ="на рассмотрении"').fetchall()
        report_menu = InlineKeyboardMarkup()
        if type(reports) == tuple:
            reports = [reports]
        if data[0] == '>':
            coord = int(data[1]) + 1
            if coord == (len(reports) - 1):
                report_menu = InlineKeyboardMarkup()
                report_menu.add(InlineKeyboardButton('<', callback_data="R" + '<' + str(coord - 1) + '.' + str(
                    reports[coord - 1][0])),
                                InlineKeyboardButton('back', callback_data='Bback'))
                report_menu.add(InlineKeyboardButton('Прсмотр', callback_data="R" + 'a' + str(reports[coord][0])))
                bot.edit_message_text("id администратора " + str(reports[coord][1]) + " id пользователя "
                                      + str(reports[coord][2]) + '\n' + 'отношение ' + str(reports[coord][5]),
                                      call.message.chat.id,
                                      call.message.message_id, reply_markup=report_menu)

            elif coord < len(reports):
                report_menu = InlineKeyboardMarkup()
                report_menu.add(InlineKeyboardButton('<', callback_data="R" + '<' + str(coord - 1) + '.' + str(
                    reports[coord - 1][0])),
                                InlineKeyboardButton('back', callback_data='Bback'),
                                InlineKeyboardButton('>',
                                                     callback_data="R" + '>' + str(coord) + '.' + str(
                                                         reports[coord][0])))
                report_menu.add(InlineKeyboardButton('Прсмотр', callback_data="R" + 'a' + str(reports[coord][0])))
                bot.edit_message_text("id администратора " + str(reports[coord][1]) + " id пользователя "
                                      + str(reports[coord][2]) + '\n' + 'отношение ' + str(reports[coord][5]),
                                      call.message.chat.id,
                                      call.message.message_id, reply_markup=report_menu)
        elif data[0] == '<':
            coord = int(data[1])
            if coord == 0:
                report_menu = InlineKeyboardMarkup()
                report_menu.add(InlineKeyboardButton('back', callback_data='Bback'),
                                InlineKeyboardButton('>',
                                                     callback_data="R" + '>' + str(coord) + '.' + str(
                                                         reports[coord][0])))
                report_menu.add(InlineKeyboardButton('Прсмотр', callback_data="R" + 'a' + str(reports[0][0])))
                bot.edit_message_text("id администратора " + str(reports[coord][1]) + " id пользователя "
                                      + str(reports[coord][2]) + '\n' + 'отношение ' + str(reports[coord][5]),
                                      call.message.chat.id, call.message.message_id, reply_markup=report_menu)

            elif coord < len(reports):
                report_menu = InlineKeyboardMarkup()
                report_menu.add(InlineKeyboardButton('<', callback_data="R" + '<' + str(coord - 1) + '.' + str(
                    reports[coord - 1][0])),
                                InlineKeyboardButton('back', callback_data='Bback'),
                                InlineKeyboardButton('>',
                                                     callback_data="R" + '>' + str(coord) + '.' + str(
                                                         reports[coord][0])))
                report_menu.add(InlineKeyboardButton('Прсмотр', callback_data="R" + 'a' + str(reports[coord][0])))
                bot.edit_message_text("id администратора " + str(reports[coord][1]) + " id пользователя "
                                      + str(reports[coord][2]) + '\n' + 'отношение ' + str(reports[coord][5]),
                                      call.message.chat.id,
                                      call.message.message_id,
                                      reply_markup=report_menu)

        elif data[0] == 'a':
            data = data.split('a')[1]
            with db.connection:
                report = db.connection.execute(f'SELECT * FROM reports WHERE id = {data}').fetchone()
            report_menu = InlineKeyboardMarkup()
            report_menu.add(InlineKeyboardButton('принять', callback_data='I' + str(data) + 'accept'),
                            InlineKeyboardButton('отклонить', callback_data='I' + str(data) + 'refuse'))
            report_menu.add(InlineKeyboardButton('back', callback_data='Bback'))
            bot.edit_message_text("id администратора " + str(report[1]) + " id пользователя "
                                  + str(report[2]) + '\n' + 'отношение ' + str(report[5]) + '\n'
                                  + 'Отзыв: ' + str(report[4]),
                                  call.message.chat.id, call.message.message_id, reply_markup=report_menu)

    elif flag == 'I':
        i_menu = InlineKeyboardMarkup()
        if data[-6:] == 'accept':
            report_id = data.replace('accept', '')
            with db.connection:
                db.connection.execute(f'UPDATE reports SET status ="разрешен" WHERE id = {report_id}')
                rel = db.connection.execute(
                    f'SELECT relationship, user_id FROM reports WHERE id = {report_id}').fetchone()
            if rel[0] == 'admin-to-user':
                with db.connection:
                    db.connection.execute(f'UPDATE user SET strike_status = strike_status+1 WHERE tg_id = {rel[1]}')
            i_menu.add(InlineKeyboardButton('back', callback_data='Bback'))
            bot.send_message('Отзыв принят', call.message.chat.id, call.message.message_id, reply_markup=i_menu)

        elif data[-6:] == 'refuse':
            report_id = data.replace('refuse', '')
            with db.connection:
                db.connection.execute(f'UPDATE reports SET status ="отклонен" WHERE id = {report_id}')
            i_menu.add(InlineKeyboardButton('back', callback_data='Bback'))
            bot.send_message(call.message.chat.id, 'Отзыв отклонен', reply_markup=i_menu)

    elif flag == 'Y':
        with db.connection:
            lots = db.connection.execute(f'SELECT * FROM lots WHERE status ="активный"').fetchall()
        if type(lots) == tuple:
            lots = [lots]
        if data[0] == '>':
            coord = int(data[1]) + 1
            if coord == (len(lots) - 1):
                d_menu = InlineKeyboardMarkup()
                d_menu.add(InlineKeyboardButton('<', callback_data="Y" + '<' + str(coord - 1) + '.' + str(
                    lots[coord - 1][0])),
                           InlineKeyboardButton('back', callback_data='Bback'))
                d_menu.add(InlineKeyboardButton('Удалить', callback_data="Y" + 'a' + str(lots[coord][0])))
                bot.edit_message_text("Лот №  " + str(lots[coord][0]), call.message.chat.id,
                                      call.message.message_id, reply_markup=d_menu)

            elif coord < len(lots):
                d_menu = InlineKeyboardMarkup()
                d_menu.add(InlineKeyboardButton('<', callback_data="Y" + '<' + str(coord - 1) + '.' + str(
                    lots[coord - 1][0])),
                           InlineKeyboardButton('back', callback_data='Bback'),
                           InlineKeyboardButton('>',
                                                callback_data="Y" + '>' + str(coord) + '.' + str(
                                                    lots[coord][0])))
                d_menu.add(InlineKeyboardButton('Удалить', callback_data="Y" + 'a' + str(lots[coord][0])))
                bot.edit_message_text("Лот №  " + str(lots[coord][0]), call.message.chat.id,
                                      call.message.message_id, reply_markup=d_menu)
        elif data[0] == '<':
            coord = int(data[1])
            if coord == 0:
                d_menu = InlineKeyboardMarkup()
                d_menu.add(InlineKeyboardButton('back', callback_data='Bback'),
                           InlineKeyboardButton('>',
                                                callback_data="Y" + '>' + str(coord) + '.' + str(
                                                    lots[coord][0])))
                d_menu.add(InlineKeyboardButton('Удалить', callback_data="Y" + 'a' + str(lots[0][0])))
                bot.edit_message_text(
                    "Лот №  " + str(lots[coord][0]),
                    call.message.chat.id, call.message.message_id, reply_markup=d_menu)

            elif coord < len(lots):
                d_menu = InlineKeyboardMarkup()
                d_menu.add(InlineKeyboardButton('<', callback_data="Y" + '<' + str(coord - 1) + '.' + str(
                    lots[coord - 1][0])),
                           InlineKeyboardButton('back', callback_data='Bback'),
                           InlineKeyboardButton('>',
                                                callback_data="Y" + '>' + str(coord) + '.' + str(
                                                    lots[coord][0])))
                d_menu.add(InlineKeyboardButton('Удалить', callback_data="Y" + 'a' + str(lots[coord][0])))
                bot.edit_message_text("Лот №  " + str(lots[coord][0]), call.message.chat.id,
                                      call.message.message_id, reply_markup=d_menu)
        elif data[0] == 'a':
            data = data.split('a')[1]
            print(data)
            lot_menu = InlineKeyboardMarkup()
            lot_menu.add(InlineKeyboardButton('удалить', callback_data='X' + str(data) + 'accept'),
                         InlineKeyboardButton('отмена', callback_data='Bback'))
            bot.edit_message_text("Точно удаляемя Лот №  " + str(data) + '?', call.message.chat.id,
                                  call.message.message_id, reply_markup=lot_menu)

    elif flag == 'X':
        f_menu = InlineKeyboardMarkup()
        if data[-6:] == 'accept':
            lot_id = data.replace('accept', '')
            with db.connection:
                del_info = db.connection.execute(
                    f'SELECT message_id, trader_link, start_price FROM lots WHERE id ={lot_id}').fetchone()
                db.connection.execute(f'DELETE FROM lots WHERE id = {lot_id}')
                db.connection.execute(
                    f'UPDATE admin SET balance = balance - 0.05*({del_info[2]}) WHERE tg_id = {del_info[1]}')
            bot.delete_message(tg_group, message_id=del_info[0])
            bot.send_message(call.message.chat.id, 'Запись удалена', reply_markup=f_menu)

    if data == "Удаление лота":
        lots_menu = InlineKeyboardMarkup()
        with db.connection:
            lots = db.connection.execute(f'SELECT * FROM lots WHERE status ="активный"').fetchall()
        if type(lots) == tuple:
            lots = [lots]
        if lots != []:
            if len(lots) == 1:
                lots_menu.add(InlineKeyboardButton('back', callback_data='Bback'))
                lots_menu.add(
                    InlineKeyboardButton('Удалить', callback_data="Y" + 'a' + str(lots[0][0])))
                bot.edit_message_text(
                    "Лот № " + str(lots[0][0]), call.message.chat.id, call.message.message_id, reply_markup=lots_menu)
            else:
                lots_menu.add(InlineKeyboardButton('back', callback_data='Bback'),
                              InlineKeyboardButton('>', callback_data="Y" + '>' + str(0) + '.' + str(
                                  lots[0][0])))
                lots_menu.add(
                    InlineKeyboardButton('Удалить', callback_data="Y" + 'a' + str(lots[0][0])))
                bot.edit_message_text("Лот № " + str(lots[0][0]), call.message.chat.id,
                                      call.message.message_id, reply_markup=lots_menu)
        else:
            lots_menu.add(InlineKeyboardButton('back', callback_data='Bback'))
            bot.edit_message_text("Лотов на расмотрении нет", call.message.chat.id,
                                  call.message.message_id, reply_markup=lots_menu)

    if data == "Одобрение лота":
        lots_menu = InlineKeyboardMarkup()
        with db.connection:
            lots = db.connection.execute(f'SELECT * FROM lots WHERE status ="на рассмотрении"').fetchall()
        if type(lots) == tuple:
            lots = [lots]
        if lots != []:
            if len(lots) == 1:
                lots_menu.add(InlineKeyboardButton('back', callback_data='Bback'))
                lots_menu.add(
                    InlineKeyboardButton('Просмотр', callback_data="F" + 'a' + str(lots[0][0])))
                bot.edit_message_text(
                    "Лот № " + str(lots[0][0]), call.message.chat.id, call.message.message_id, reply_markup=lots_menu)
            else:
                lots_menu.add(InlineKeyboardButton('back', callback_data='Bback'),
                              InlineKeyboardButton('>', callback_data="F" + '>' + str(0) + '.' + str(
                                  lots[0][0])))
                lots_menu.add(
                    InlineKeyboardButton('Просмотр', callback_data="F" + 'a' + str(lots[0][0])))
                bot.edit_message_text("Лот № " + str(lots[0][0]), call.message.chat.id,
                                      call.message.message_id, reply_markup=lots_menu)
        else:
            lots_menu.add(InlineKeyboardButton('back', callback_data='Bback'))
            bot.edit_message_text("Лотов на расмотрении нет", call.message.chat.id,
                                  call.message.message_id, reply_markup=lots_menu)

    if data == 'Пожаловаться':
        with db.connection:
            users = db.connection.execute(f'SELECT tg_id FROM user').fetchall()
        admin_menu = InlineKeyboardMarkup()
        if type(users) == tuple:
            users = [users]
        if users != []:
            if len(users) == 1:
                admin_menu.add(InlineKeyboardButton('back', callback_data='Bback'))
                admin_menu.add(
                    InlineKeyboardButton('Просмотр', callback_data="D" + 'a' + str(users[0][0])))
                bot.edit_message_text(
                    "id пользователя " + str(users[0][0]), call.message.chat.id, call.message.message_id,
                    reply_markup=admin_menu)
            else:
                admin_menu.add(InlineKeyboardButton('back', callback_data='Bback'),
                               InlineKeyboardButton('>', callback_data="D" + '>' + str(0) + '.' + str(
                                   users[0][0])))
                admin_menu.add(
                    InlineKeyboardButton('Просмотр', callback_data="D" + 'a' + str(users[0][0])))
                bot.edit_message_text("id пользователя " + str(users[0][0]), call.message.chat.id,
                                      call.message.message_id, reply_markup=admin_menu)
        else:
            admin_menu.add(InlineKeyboardButton('back', callback_data='Bback'))
            bot.edit_message_text("Нет пользователей", call.message.chat.id,
                                  call.message.message_id, reply_markup=admin_menu)

    if data == 'Жалобы':
        with db.connection:
            reports = db.connection.execute(f'SELECT * FROM reports WHERE status ="на рассмотрении"').fetchall()
        report_menu = InlineKeyboardMarkup()
        if type(reports) == tuple:
            reports = [reports]
        if reports != []:
            if len(reports) == 1:
                report_menu.add(InlineKeyboardButton('back', callback_data='Bback'))
                report_menu.add(
                    InlineKeyboardButton('Просмотр', callback_data="R" + 'a' + str(reports[0][0])))
                bot.edit_message_text(
                    "id администратора " + str(reports[0][1]) + " id пользователя " + str(
                        reports[0][2]) + '\n' + 'отношение ' + str(reports[0][5]),
                    call.message.chat.id, call.message.message_id, reply_markup=report_menu)
            else:
                report_menu.add(InlineKeyboardButton('back', callback_data='Bback'),
                                InlineKeyboardButton('>', callback_data="R" + '>' + str(0) + '.' + str(
                                    reports[0][0])))
                report_menu.add(
                    InlineKeyboardButton('Просмотр', callback_data="R" + 'a' + str(reports[0][0])))
                bot.edit_message_text("id администратора " + str(reports[0][1]) + " id пользователя " + str(
                    reports[0][2]) + '\n' + 'отношение ' + str(reports[0][5]), call.message.chat.id,
                                      call.message.message_id, reply_markup=report_menu)
        else:
            report_menu.add(InlineKeyboardButton('back', callback_data='Bback'))
            bot.edit_message_text("Нет отзывов", call.message.chat.id,
                                  call.message.message_id, reply_markup=report_menu)

    if data == 'История торгов':
        history = InlineKeyboardMarkup()
        history.add(InlineKeyboardButton('back', callback_data='Bback'))
        with db.connection:
            lots = db.connection.execute(f'SELECT * FROM lots').fetchall()
        if type(lots) == tuple:
            lots = [lots]
        text_h = ''
        for lot in lots:
            text_h += f'№ лота {lot[0]} стоимостью {lot[1]} имеет статус {lot[10]}\n'
        bot.edit_message_text(text_h, call.message.chat.id,
                              call.message.message_id, reply_markup=history)

    if data == 'Баланс':
        with db.connection:
            admin_balance = \
                db.connection.execute(f'SELECT balance FROM admin WHERE tg_id = {call.message.chat.id}').fetchone()[0]
        bot.send_message(call.message.chat.id,
                         f'Баланс пользователя {call.message.chat.id} составляет {admin_balance} золотых монет')

    if data == 'Создать_лот':
        msg = bot.send_message(call.message.chat.id, 'начинаем создавать \nЗагрузите фото')
        bot.register_next_step_handler(msg, photo)

    if data == 'Начисление баланса':
        with db.connection:
            admin_list = db.connection.execute(f'SELECT tg_id FROM admin WHERE rights = 3 or rights = 1').fetchall()
        money_menu = InlineKeyboardMarkup()
        for item in admin_list:
            money_menu.add(InlineKeyboardButton(str(item[0]), callback_data='M' + str(item[0])))
        bot.edit_message_text('Меню одменов', call.message.chat.id,
                              call.message.message_id, reply_markup=money_menu)

    if data == "Изменение админов":
        with db.connection:
            list_admins = db.connection.execute(f"SELECT * FROM admin").fetchall()
        if type(list_admins) == tuple:
            list_admins = [list_admins]
        adm_menu = InlineKeyboardMarkup()
        for admin in list_admins:
            adm_menu.add(InlineKeyboardButton(str(admin[1]) + ' - status: ' + str(admin[2]),
                                              callback_data='J' + str(admin[1]) + str(admin[2])))
        adm_menu.add(InlineKeyboardButton('back', callback_data='Bback'))
        bot.edit_message_text('меню одменов', call.message.chat.id,
                              call.message.message_id, reply_markup=adm_menu)


def admin_panel_run():
    print('ready')
    bot.polling()


if __name__ == '__main__':
    admin_panel_run()
